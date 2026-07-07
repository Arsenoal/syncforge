#!/usr/bin/env python3
"""Generate SyncForge roadmap DOCX files (1.0.0 through 2.0.0 + 1.0 P0 checklist)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_1_0_TO_2_0 = ROOT / "docs" / "SyncForge-Roadmap-1.0-to-2.0.docx"
OUTPUT_P0 = ROOT / "docs" / "SyncForge-1.0-P0.docx"

DOC_DATE = "7 July 2026"
BASELINE = "1.0.0 GA"
TARGET_1_1 = "1.1.0"


def _style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)


def _shade_cell(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        _shade_cell(hdr[i], "D5E8F0")
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, text in enumerate(row):
            cells[ci].text = text
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _job_table(doc: Document, jobs: list[tuple[str, str, str, str]]) -> None:
    _add_table(doc, ["ID", "Job", "Priority", "Area"], jobs, [0.9, 3.8, 0.7, 0.9])


def build_1_0_to_2_0() -> None:
    doc = Document()
    _style_doc(doc)

    title = doc.add_heading("SyncForge Roadmap", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Versions 1.0.0 through 2.0.0")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(60, 60, 60)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run(f"Document date: {DOC_DATE}\n").italic = True
    meta.add_run(f"Baseline: {BASELINE}\n").italic = True
    meta.add_run("Maven group: studio.syncforge · Packages: dev.syncforge.*").italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "SyncForge 1.0 establishes a semver-stable Android and common sync contract: "
        "outbox → push → pull → configurable conflict strategies. Versions 1.1–1.5 deepen "
        "developer experience, conflict tooling, platform parity, ecosystem adapters, and "
        "observability without changing the core sync loop. Version 2.0 reserves opt-in "
        "architectural extensions (field-level CRDT strategies, optional op-log sync mode, "
        "KMP platform graduation) that may introduce breaking API or REST contract changes."
    )

    doc.add_heading("Release arc", level=2)
    _add_table(
        doc,
        ["Version", "Codename", "Headline"],
        [
            ("1.0.0", "Stable", "API freeze, Maven Central 1.0, remove pre-1.0 deprecations"),
            ("1.1.x", "Wire-up", "EntityStore + injectable Ktor client, DI, auth hardening, cursor"),
            ("1.2.x", "Merge-smart", "Per-entity strategies, gitLike { }, CRDT, KSP field-merge"),
            ("1.3.x", "Everywhere", "Desktop sample, iOS SPM/XCFramework, CMP debug UI"),
            ("1.4.x", "Ecosystem", "Spring/GraphQL/Supabase transports, multi-device E2E, version catalog"),
            ("1.5.x", "Operate", "OpenTelemetry, SyncHealth dashboard, hierarchical recipes"),
            ("2.0.0", "Converge", "Optional CRDT/op-log channel, stable KMP, REST v2 (if needed)"),
        ],
        [0.75, 1.0, 4.5],
    )

    doc.add_heading("2. Strategic themes", level=1)
    _add_table(
        doc,
        ["Theme", "1.0", "1.x", "2.0"],
        [
            ("Core sync loop", "Stable", "Hardening only", "Optional op-log / CRDT doc mode"),
            ("Conflict resolution", "Per-entity conflicts { }", "gitLike three-way, strategy catalog", "crdt { } first-class"),
            ("Android", "Primary stable target", "DI modules, ProGuard sign-off", "Legacy cleanup complete"),
            ("iOS / desktop / macOS", "Experimental DSLs", "Sample parity, SPM binary", "Graduate to stable"),
            ("Backend / transport", "REST v1 (KtorSyncTransport)", "GraphQL, Supabase, Spring adapters", "REST v2 if needed; SyncTransport pluggable"),
            ("Distribution", "BOM + Gradle plugin", "Version catalog, DI artifacts", "SPM + Maven parity"),
        ],
        [1.4, 1.5, 1.5, 1.9],
    )

    doc.add_page_break()
    doc.add_heading("3. Version 1.0.0 — First stable release", level=1)
    doc.add_paragraph(
        "Goal: Ship a trustworthy 1.0 — documented, tested, Maven Central, semver guarantees "
        "on the Android-primary API."
    )

    doc.add_heading("3.1 Already complete (0.6 → 1.0.0)", level=2)
    _bullets(
        doc,
        [
            "Sample App Proof: tasks + notes + tags, multi-screen, per-entity conflicts { }",
            "iOS parity: sample-ios-shared + SwiftUI TabView, MultiEntityUITests",
            "SQLDelight default + RoomToSqlDelightMigrator + sign-off tests",
            "RefreshingSyncAuthProvider, 401 retry, WorkManager + iOS background sync",
            "syncforge-server, backend-starter, mock-server, consumer-smoke",
            "Maven Central BOM, studio.syncforge.android Gradle plugin",
            "CI: androidE2e, iosE2e, verifyReleaseSignOff, verifyConsumerSmoke",
            "ConflictResolver family removed; useSqlDelightPersistence() removed",
        ],
    )

    doc.add_heading("3.2 P0 checklist (complete)", level=2)
    _job_table(
        doc,
        [
            ("1.0-P0-01", "API graduation — remove @ExperimentalSyncForgeApi from stable surfaces", "P0", "API ✅"),
            ("1.0-P0-02", "Remove useRoomPersistence() — document upgrade path", "P0", "Android ✅"),
            ("1.0-P0-03", "Publish 1.0.0 to Maven Central (all artifacts)", "P0", "Dist ✅"),
            ("1.0-P0-04", "1.0 sign-off matrix — tag v1.0.0", "P0", "QA ✅"),
            ("1.0-P0-05", "macOS publish — iOS/macOS KMP targets on Central", "P0", "CI ✅"),
            ("1.0-P0-06", "Docs freeze — CHANGELOG, MODULES, GETTING_STARTED", "P0", "Docs ✅"),
        ],
    )

    doc.add_heading("3.3 P1 (may slip to 1.0.1)", level=2)
    _job_table(
        doc,
        [
            ("1.0-P1-01", "Upgrade guide: pre-0.6 Room → 0.6+ SQLDelight", "P1", "Docs"),
            ("1.0-P1-02", "Conflict recipes for 3+ entity types in RECIPES.md", "P1", "Docs"),
            ("1.0-P1-03", "Consumer ProGuard/R8 rules documented + tested", "P1", "Android"),
            ("1.0-P1-04", "Integration tests: retry exhaustion, multi-page pull", "P1", "QA"),
            ("1.0-P1-05", "Performance test: 1000+ outbox entries", "P1", "QA"),
            ("1.0-P1-06", "Security doc pass: TLS, token storage, no secrets in logs", "P1", "Docs"),
            ("1.0-P1-07", "External dogfood or third-party integration", "P1", "Community"),
        ],
    )

    doc.add_page_break()
    doc.add_heading("4. Version 1.1.x — Integration & persistence", level=1)
    doc.add_paragraph(
        "Goal: EntityStore abstraction (any app database), injectable Ktor HttpClient, "
        "DI helpers, auth graduation, encrypted TokenStore, CharArray credential APIs, "
        "DataStore cursor. Core sync unchanged."
    )
    _job_table(
        doc,
        [
            ("1.1-01", "DI recipes — Koin + Hilt in RECIPES.md", "P0", "DX"),
            ("1.1-02", "syncforge-integration-koin optional artifact", "P1", "DX"),
            ("1.1-03", "syncforge-integration-hilt optional artifact", "P1", "DX"),
            ("1.1-04", "Graduate built-in auth DSL to stable", "P1", "API"),
            ("1.1-05", "DataStore Preferences cursor (KMP)", "P1", "Persistence"),
            ("1.1-06", "Sample Hilt or Koin variant", "P2", "Sample"),
            ("1.1-07", "SyncForgeBuilder graduation", "P2", "API"),
            ("1.1-09", "EntityStore contract in commonMain", "P0", "API"),
            ("1.1-10", "KSP @SyncForgeStore — handlers from any EntityStore", "P0", "KSP"),
            ("1.1-11", "Optional store adapters (room, realm, in-memory)", "P1", "DX"),
            ("1.1-12", "Docs + plugin: Room optional; BYO store path", "P1", "Docs"),
            ("1.1-13", "SyncHttpClient contract — Ktor-backed REST executor", "P0", "Network"),
            ("1.1-14", "RestSyncTransport — push/pull via SyncHttpClient", "P0", "Network"),
            ("1.1-15", "syncforge-network-ktor adapter (extract from core)", "P1", "Network"),
            ("1.1-16", "DSL httpClient { } — inject app HttpClient", "P1", "DX"),
            ("1.1-17", "Encrypted TokenStore — Keychain / EncryptedSharedPreferences", "P1", "Security"),
            ("1.1-18", "login/register CharArray overloads + wipe semantics", "P1", "Security"),
        ],
    )
    doc.add_paragraph(
        "Auth security: passwords never persisted; tokens encrypted at rest (1.1-17); "
        "optional CharArray credential APIs with documented wipe (1.1-18). "
        "External IdP (Firebase Auth) remains recommended for production."
    )
    doc.add_paragraph(
        "Network: Ktor-only SyncHttpClient executes REST; RestSyncTransport maps "
        "/sync/push and /sync/pull. User may inject app HttpClient; SyncForge owns route + DTO mapping."
    )
    doc.add_paragraph(
        "Entity store: App data via EntityStore / EntitySyncHandler (Room, Realm, or custom). "
        "SyncForge outbox/conflicts stay in SQLDelight. Optional syncforge-store-* adapters; "
        ":syncforge core has no Room/Realm dependency."
    )
    doc.add_paragraph(
        "DI architecture: :syncforge has no Koin/Dagger dependency. Optional "
        "syncforge-integration-koin and syncforge-integration-hilt artifacts. App always "
        "supplies baseUrl, EntityRegistry (handlers or stores), and conflicts { }."
    )

    doc.add_heading(f"4.1 {TARGET_1_1} GitHub issues breakdown", level=2)
    doc.add_paragraph(
        f"Milestone {TARGET_1_1} (Wire-up): P0 network + entity-store in parallel, then "
        "security, DI docs, release gate. Full issue tables: docs/SyncForge-1.1-Issues.docx."
    )
    _add_table(
        doc,
        ["Epic", "Issues", "P0 jobs"],
        [
            ("A — Injectable HTTP", "A1→A5", "1.1-13, 1.1-14, 1.1-16"),
            ("B — EntityStore", "B1→B5", "1.1-09, 1.1-10, 1.1-11, 1.1-12"),
            ("C — Auth security", "C1→C3", "1.1-17, 1.1-18, 1.1-04"),
            ("D — DI & DX", "D1→D4", "1.1-01 (P0 docs); 1.1-02/03 P1"),
            ("E — Cursor", "E1", "1.1-05"),
            ("F — Release", "F1→F4", "BOM, acceptance, tag v1.1.0"),
        ],
        [1.5, 1.0, 2.8],
    )
    doc.add_paragraph("Suggested order: Weeks 1–2 A+B foundations → Weeks 3–4 adapters + cursor → Week 5 security + docs → Week 6 release.")

    doc.add_heading("5. Version 1.2.x — Conflict evolution (per-entity + git-like)", level=1)
    doc.add_paragraph(
        "Goal: Per-entity conflict resolver from a strategy catalog (accept-remote, merge, "
        "gitLike three-way, deferToUser, …). App can pick strategy per entity type, including "
        "runtime overrides. CRDT + KSP field-merge — not a replacement for outbox → push → pull."
    )
    _job_table(
        doc,
        [
            ("1.2-01", "CRDT primitives: LwwRegister, OrSet, GCounter", "P0", "Conflict"),
            ("1.2-02", "crdt { } conflict strategy (experimental)", "P0", "Conflict"),
            ("1.2-03", "KSP field-merge annotations @Lww, @OrSet, @GCounter", "P1", "KSP"),
            ("1.2-04", "Tombstone-aware merge recipes", "P1", "Docs"),
            ("1.2-05", "Multi-device E2E (concurrent edit)", "P1", "QA"),
            ("1.2-06", "CONFLICT_RESOLUTION.md v2 — per-entity matrix", "P1", "Docs"),
            ("1.2-07", "Merge-base snapshot store (three-way merge)", "P0", "Conflict"),
            ("1.2-08", "gitLike { } — threeWayMerge + deferToUser fallback", "P0", "Conflict"),
            ("1.2-09", "ConflictStrategyKind catalog + fromKind()", "P0", "API"),
            ("1.2-10", "Runtime updateConflictPolicy() from app prefs", "P1", "API"),
            ("1.2-11", "Outbox reconcile on AcceptRemote / Custom merge", "P1", "Sync"),
        ],
    )
    _add_table(
        doc,
        ["Entity (example)", "Recommended strategy"],
        [
            ("Notes (:sample)", "alwaysRemote() / ACCEPT_REMOTE"),
            ("Tasks (:sample)", "merge { } or gitLike { }"),
            ("Settings / config", "alwaysRemote()"),
            ("Tags, collaborators", "crdt { field(\"tags\") { orSet() } }"),
            ("Counters", "crdt { field(\"views\") { gCounter() } }"),
            ("Legal / high-value", "deferToUser() or gitLike fallback"),
            ("Low-stakes rows", "lastWriteWins()"),
        ],
        [2.2, 3.1],
    )

    doc.add_page_break()
    doc.add_heading("6. Version 1.3.x — Platform parity", level=1)
    _job_table(
        doc,
        [
            ("1.3-01", ":sample-desktop minimal CLI or Compose", "P0", "Desktop"),
            ("1.3-02", "Graduate SyncForge.ios { } to stable", "P1", "iOS"),
            ("1.3-03", "Graduate desktop/macos DSLs", "P1", "Desktop"),
            ("1.3-04", "Swift Package Manager / XCFramework publish", "P1", "iOS"),
            ("1.3-05", "Compose Multiplatform conflict/debug UI", "P2", "UI"),
            ("1.3-06", "Shake-to-open SyncDebugLauncher", "P2", "DX"),
            ("1.3-07", "SKIE Swift API review + documentation", "P1", "iOS"),
        ],
    )

    doc.add_heading("7. Version 1.4.x — Ecosystem", level=1)
    doc.add_paragraph(
        "Goal: General transport control via SyncDeltaStore + DeltaStoreSyncTransport "
        "(Firebase, Supabase, custom BaaS). GraphQL and REST adapters share the same "
        "push/pull semantics. SyncTransport is the engine boundary; REST remains 1.0 default."
    )
    _job_table(
        doc,
        [
            ("1.4-01", "Spring Boot backend starter", "P0", "Ecosystem"),
            ("1.4-02", "SyncDeltaStore port + DeltaStoreSyncTransport (transport-core)", "P0", "Transport"),
            ("1.4-03", "syncforge-transport-supabase — SyncDeltaStore impl", "P1", "Transport"),
            ("1.4-04", "syncforge-transport-firebase — SyncDeltaStore impl", "P1", "Transport"),
            ("1.4-05", "Gradle version catalog for consumers", "P1", "Dist"),
            ("1.4-06", "Multi-device E2E — two emulators", "P1", "QA"),
            ("1.4-07", "Backend contract test kit (REST + SyncDeltaStore)", "P2", "QA"),
            ("1.4-08", "syncforge-transport-graphql client adapter", "P1", "Transport"),
            ("1.4-09", "GraphQL schema + resolver recipes", "P1", "Docs"),
            ("1.4-10", "Custom transport guide — BYO SyncTransport or SyncDeltaStore", "P2", "Docs"),
        ],
    )

    doc.add_heading("8. Version 1.5.x — Observability", level=1)
    _job_table(
        doc,
        [
            ("1.5-01", "Structured tracing / OpenTelemetry hooks", "P0", "Observability"),
            ("1.5-02", "SyncHealth metrics expansion", "P1", "Observability"),
            ("1.5-03", "Full SyncHealth dashboard UI", "P1", "DX"),
            ("1.5-04", "Hierarchical sync recipes", "P1", "Docs"),
            ("1.5-05", "Rate limiting + backoff policies", "P2", "Network"),
            ("1.5-06", "Audit log export (CSV/JSON)", "P2", "DX"),
        ],
    )

    doc.add_page_break()
    doc.add_heading("9. Version 2.0.0 — Major release vision", level=1)
    doc.add_paragraph(
        "Optional second sync mode for CRDT-heavy products while keeping REST entity sync as default."
    )
    _add_table(
        doc,
        ["Theme", "Description", "Breaking?"],
        [
            ("Stable KMP everywhere", "All platform DSLs stable", "Maybe"),
            ("crdt { } stable", "Field CRDT strategy graduates", "No"),
            ("Op-log / CRDT document channel", "New :syncforge-crdt module", "Yes"),
            ("REST API v2", "Only if op-log needs new endpoints", "Yes"),
            ("KSP-generated DI modules", "Koin/Hilt stubs from @SyncForgeEntity", "No"),
        ],
        [1.6, 2.8, 0.9],
    )

    doc.add_heading("2.0 explicit non-goals", level=2)
    _bullets(
        doc,
        [
            "Replacing Room/SQLDelight as the app entity store",
            "Full real-time WebSocket sync as the only mode",
            "Automatic CRDT for whole JSON blobs without schema",
            "Bundling Koin or Dagger into core :syncforge",
        ],
    )

    doc.add_heading("10. REST API evolution", level=1)
    _add_table(
        doc,
        ["Library version", "REST contract", "Notes"],
        [
            ("1.0.x", "v1 frozen", "POST /sync/push, GET /sync/pull per REST_API.md"),
            ("1.x", "v1 + auth routes", "/auth/* with built-in auth"),
            ("2.0", "v1 + optional v2", "v2 only if op-log channel needs it"),
        ],
        [1.0, 1.2, 3.1],
    )

    doc.add_heading("11. Risk register", level=1)
    _add_table(
        doc,
        ["Risk", "Severity", "Mitigation"],
        [
            ("1.0 API surface too large", "Medium", "Stable = Android + common; KMP experimental until 1.3"),
            ("CRDT wrong merges for deletes", "High", "Keep deferToUser(); document tombstones"),
            ("DI artifact fragmentation", "Low", "Max two optional modules; recipes first"),
            ("iOS SPM publish complexity", "Medium", "XCFramework from macOS CI; KMP fallback"),
            ("REST v2 splits ecosystem", "High", "Defer to 2.0; v1 long support window"),
            ("Multi-device E2E flakiness", "Medium", "Nightly only; mock-server health gate"),
        ],
        [2.0, 0.8, 2.5],
    )

    doc.add_heading("12. Indicative timeline (from 1.0.0 GA)", level=1)
    doc.add_paragraph("Assuming part-time maintenance or a small team:")
    _add_table(
        doc,
        ["Phase", "Duration", "Cumulative"],
        [
            ("1.0.0 soak → GA", "2–4 weeks", "~1 month"),
            ("1.1.x", "6–8 weeks", "~3 months"),
            ("1.2.x", "8–10 weeks", "~5 months"),
            ("1.3.x", "8–10 weeks", "~8 months"),
            ("1.4.x", "6–8 weeks", "~10 months"),
            ("1.5.x", "6–8 weeks", "~12 months"),
            ("2.0.0 RC + GA", "8–12 weeks", "~15 months"),
        ],
        [1.8, 1.2, 1.2],
    )

    doc.add_page_break()
    doc.add_heading("13. Acceptance criteria — 1.0.0 sign-off (complete)", level=1)
    _bullets(
        doc,
        [
            "All 1.0-P0 jobs completed and verified — GA July 2026, tag v1.0.0.",
            "Maven Central 1.0.0 + consumer-smoke from Central only.",
            "Sample App Proof E2E green; stable API surfaces without @ExperimentalSyncForgeApi.",
            "External dogfood (P1-07) remains optional community follow-up.",
        ],
    )

    doc.add_heading("13.1 Acceptance criteria — 1.1.0 (target)", level=1)
    _bullets(
        doc,
        [
            "SyncHttpClient + RestSyncTransport; KtorSyncTransport backward compatible.",
            "httpClient { } DSL + documented injectable HttpClient sample.",
            "EntityStore + @SyncForgeStore KSP; Room adapter optional module.",
            "Encrypted TokenStore + CharArray auth overloads.",
            "DataStore cursor (Android); RECIPES DI section.",
            "BOM optional artifacts; no 1.0 stable API breaks.",
        ],
    )

    doc.add_heading("14. Acceptance criteria — 2.0.0 sign-off", level=1)
    _bullets(
        doc,
        [
            "All P0 jobs for 1.5–2.0 completed.",
            "Entity sync (1.x path) unchanged for consumers not opting into new modules.",
            "Platform DSLs stable on Android, iOS, desktop, macOS.",
            "crdt { } graduated or explicit 2.x plan documented.",
            "DI integration artifacts published and in GETTING_STARTED.",
            "Spring starter + one hosted adapter documented.",
            "Tracing hooks documented with sample exporter.",
            "REST versioning policy updated; v1 deprecation timeline if v2 ships.",
            "Full upgrade guide 1.x → 2.0 with breaking changes enumerated.",
            "Security review on auth, token storage, transport defaults.",
        ],
    )

    doc.add_heading("15. Related documents", level=1)
    _add_table(
        doc,
        ["Document", "Purpose"],
        [
            ("docs/ROADMAP.md", "Pre-1.0 phases and current status"),
            ("docs/ROADMAP_1_0_TO_2_0.md", "Markdown source for this document"),
            ("docs/MODULES.md", "API stability by area"),
            ("docs/REST_API.md", "Backend contract + versioning"),
            ("docs/SyncForge-1.0-P0.docx", "Post-1.0 P1 / community checklist"),
            ("docs/SyncForge-1.1-Issues.docx", "1.1.0 GitHub issues breakdown"),
        ],
        [2.5, 3.0],
    )

    OUTPUT_1_0_TO_2_0.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_1_0_TO_2_0)
    print(f"Wrote {OUTPUT_1_0_TO_2_0}")


def build_p0_checklist() -> None:
    doc = Document()
    _style_doc(doc)

    title = doc.add_heading("SyncForge 1.0.0 — P1 Follow-ups", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Post-1.0.0 P1 and community items (P0 release gate complete).")
    doc.add_paragraph(f"Date: {DOC_DATE}")
    doc.add_paragraph(f"Baseline: {BASELINE}")
    doc.add_paragraph()

    sections = [
        (
            "Documentation (P1)",
            [
                "1.0-P1-01: Upgrade guide — pre-0.6 Room → SQLDelight",
                "1.0-P1-02: Conflict recipes for 3+ entity types in RECIPES.md",
                "1.0-P1-06: Security doc pass — TLS, token storage, no secrets in logs",
            ],
        ),
        (
            "Android & QA (P1)",
            [
                "1.0-P1-03: Consumer ProGuard/R8 rules documented + tested",
                "1.0-P1-05: Performance test — 1000+ outbox entries, batch push",
            ],
        ),
        (
            "Community (P1)",
            [
                "1.0-P1-07: External dogfood or documented third-party integration",
            ],
        ),
        (
            "Next milestone — 1.1.0 Wire-up",
            [
                "See docs/SyncForge-1.1-Issues.docx for GitHub issue breakdown",
                "Epic A: SyncHttpClient + RestSyncTransport + httpClient { } DSL",
                "Epic B: EntityStore + @SyncForgeStore KSP + optional store adapters",
                "Epic C–F: Encrypted tokens, DI recipes, DataStore cursor, release gate",
            ],
        ),
    ]

    doc.add_heading("1.0.0 status", level=1)
    doc.add_paragraph(
        "GA on Maven Central (studio.syncforge:1.0.0) with tag v1.0.0. "
        "Items below are follow-ups, not release blockers. Next: milestone 1.1.0."
    )

    for heading, items in sections:
        doc.add_heading(heading, level=2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    doc.save(OUTPUT_P0)
    print(f"Wrote {OUTPUT_P0}")


def main() -> None:
    build_1_0_to_2_0()
    build_p0_checklist()


if __name__ == "__main__":
    main()
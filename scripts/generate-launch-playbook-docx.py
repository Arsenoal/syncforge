#!/usr/bin/env python3
"""Generate SyncForge GitHub launch & 1.0 release playbook (DOCX)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SyncForge-GitHub-Launch-Playbook.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D5E8F0") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, text in enumerate(row):
            cells[c_idx].text = text
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # Title page
    title = doc.add_heading("SyncForge", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("GitHub Launch & 1.0 Release Playbook")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(30, 70, 120)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Repository: github.com/Arsenoal/syncforge\n").italic = True
    meta.add_run("Document date: 2 July 2026\n").italic = True
    meta.add_run("Current baseline: 0.9.0-rc.1 (tagged)\n").italic = True
    meta.add_run("Target stable release: 1.0.0").italic = True

    doc.add_page_break()

    # Table of contents placeholder (manual — Word can refresh TOC on open)
    doc.add_heading("Contents", level=1)
    toc_items = [
        "1. Executive summary",
        "2. Current state",
        "3. Release strategy (RC soak → 1.0.0)",
        "4. 1.0.0 scope",
        "5. API stability at 1.0",
        "6. Platform support for developers",
        "7. Maven Central checklist",
        "8. GitHub & growth playbook",
        "9. Launch content kit",
        "10. Timeline",
        "11. Post-1.0 backlog (1.1+)",
        "12. Decisions we are not taking for 1.0",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Executive summary
    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "SyncForge is an offline-first sync library for Android and Kotlin Multiplatform. "
        "Phase 6 (KMP persistence, SQLDelight default, iOS/desktop/macOS targets, E2E tests) "
        "is complete. Distribution plumbing (Maven BOM, Gradle plugin, consumer smoke, "
        "publish-release workflow) is in place."
    )
    doc.add_paragraph(
        "The next milestone is not a feature-heavy release — it is 1.0.0 as a stabilization "
        "release: Maven Central, semver-stable public API, and a focused GitHub launch to "
        "attract early adopters."
    )
    doc.add_paragraph("Recommended path:")
    bullets(
        doc,
        [
            "Soak 0.9.0-rc.1 for one to two weeks on Maven Central.",
            "Fix blockers → tag 0.9.0-rc.2 if needed; otherwise ship 1.0.0.",
            "Launch on GitHub when 1.0.0 is on Central — not before.",
            "Defer pure Swift SDK, DataStore cursor, Supabase adapter, and desktop sample to 1.1+.",
        ],
    )

    # 2. Current state
    doc.add_heading("2. Current state", level=1)
    doc.add_heading("Completed", level=2)
    bullets(
        doc,
        [
            "verifyReleaseSignOff passes (JVM + Android unit tests, server tests, consumer smoke).",
            "Tag v0.9.0-rc.1 pushed; Publish Release workflow configured (macos-latest, all KMP targets).",
            "Android: SyncForge.android { }, KSP codegen, SQLDelight outbox, Room migrator, WorkManager sample.",
            "iOS: SyncForge.ios { }, SwiftUI sample (ios-sample/), XCUITest E2E, BGTaskScheduler background sync.",
            "KMP: JVM desktop + native macOS DSLs, SKIE Swift interop.",
            "Backend: :syncforge-server contract, :backend-starter with auth routes.",
            "Experimental built-in auth DSL (auth { }, register/login/logout, TokenStore).",
            "License: Apache 2.0.",
        ],
    )
    doc.add_heading("In progress / pending", level=2)
    bullets(
        doc,
        [
            "Maven Central: close and release staging repo after CI publish succeeds.",
            "Post-publish verification: consumer-smoke compiles from mavenCentral() only.",
            "README / ROADMAP version strings still reference 0.6.0-SNAPSHOT in places — update for 1.0.",
            "Documented CocoaPods / SPM consumer integration (deferred to 1.1).",
        ],
    )

    # 3. Release strategy
    doc.add_heading("3. Release strategy (RC soak → 1.0.0)", level=1)
    doc.add_paragraph(
        "Yes — waiting one to two weeks on 0.9.0-rc.1 before tagging 1.0.0 is the right approach. "
        "This is standard practice for a first public release candidate, not unnecessary delay."
    )
    doc.add_heading("What the soak period is for", level=2)
    bullets(
        doc,
        [
            "Confirm Maven Central artifacts resolve (BOM, plugin, syncforge, ksp, persistence).",
            "Run consumer-smoke/android-minimal against Central only (remove mavenLocal).",
            "Dogfood in a real app or the :sample / ios-sample apps on Android + iOS.",
            "Collect issues from early adopters; fix P0 bugs before 1.0.",
        ],
    )
    doc.add_heading("Release decision tree", level=2)
    add_table(
        doc,
        ["Situation", "Action"],
        [
            ["Central publish + consumer smoke OK, no P0 bugs", "Tag v1.0.0 and publish"],
            ["Critical bug or broken Central artifacts", "Fix, tag v0.9.0-rc.2, re-soak briefly"],
            ["Auth or migrator regression found", "Fix in rc.2; do not ship 1.0 with known data-loss risk"],
            ["Only doc / README friction", "Fix in 1.0.0 or 1.0.1 — not a reason to delay weeks"],
        ],
    )
    doc.add_paragraph(
        "Ship 1.0.0 when: Central is verified, no P0 bugs, and you are comfortable declaring API stability — "
        "not on a fixed calendar date."
    )

    # 4. 1.0.0 scope
    doc.add_heading("4. 1.0.0 scope", level=1)
    doc.add_heading("Must-have (blocking 1.0.0)", level=2)
    bullets(
        doc,
        [
            "Finish Maven Central distribution: CI publish, Sonatype close/release, Central verification.",
            "API stability sign-off — graduate platform DSLs; document stable vs experimental surfaces.",
            "Documentation polish: GETTING_STARTED coordinates → 1.0.0; fix README/ROADMAP version drift.",
            "Short Upgrading from 0.9.x guide (Room migrator, removed APIs).",
            "iOS consumer guide: Gradle framework + Xcode (SPM can wait for 1.1).",
            "Run verifyReleaseSignOff on release branch; tag v1.0.0; publish and verify Central.",
        ],
    )
    doc.add_heading("Should-have (if time before 1.0)", level=2)
    bullets(
        doc,
        [
            "Auth E2E test: Android instrumented test — login → push against :backend-starter with Bearer.",
            ":sample login/register Compose screens wired to backend-starter auth.",
            "Spring backend starter (Phase 7 item; Ktor backend-starter already exists).",
            "Remove or hard-deprecate useRoomPersistence() — migrator + sign-off tests are done.",
            "EncryptedSharedPreferences for Android TokenStore before auth graduates from experimental.",
        ],
    )
    doc.add_heading("Defer to 1.1+ (do not block 1.0)", level=2)
    add_table(
        doc,
        ["Item", "Reason to defer"],
        [
            ["DataStore multiplatform cursor", "UserDefaults / file / SharedPreferences work today"],
            [":sample-desktop packaged app", "Nice demo, not required for stable API"],
            ["Supabase adapter", "Optional transport — post-1.0"],
            ["SPM / CocoaPods packaging", "Gradle framework sufficient for 1.0 KMP users"],
            ["Pure Swift SDK (no Kotlin)", "Large rewrite — separate track, not 1.0"],
            ["Shake-to-open debug console", "DX polish"],
            ["Full sync health metrics dashboard", "Basic SyncHealth exists"],
        ],
    )

    # 5. API stability
    doc.add_heading("5. API stability at 1.0", level=1)
    doc.add_paragraph(
        "Per docs/MODULES.md, everything marked Experimental should graduate to Stable or remain "
        "explicitly experimental under semver at 1.0."
    )
    add_table(
        doc,
        ["Area", "1.0 recommendation"],
        [
            ["SyncForge.android { }, SyncManager core", "Stable — primary contract"],
            ["SyncForge.ios / .desktop / .macos DSLs", "Graduate to Stable (E2E coverage on Android + iOS)"],
            ["auth { } DSL, SyncForgeAuthService", "Keep Experimental — require @OptIn"],
            ["SyncDebug*, SyncHealth, SyncEvent", "Keep Experimental"],
            ["SyncForge.create(), builder { }", "Keep Experimental (low-level escape hatch)"],
            ["useRoomPersistence()", "Remove or hard-deprecate"],
            ["databaseName(), conflict APIs, Compose status UI", "Stable"],
        ],
    )

    # 6. Platform support
    doc.add_heading("6. Platform support for developers", level=1)
    doc.add_paragraph("After Maven Central publish, developers can integrate as follows:")
    add_table(
        doc,
        ["Project type", "Supported?", "How"],
        [
            ["Pure Android", "Yes — primary path", "dev.syncforge.android plugin + syncforge from Central"],
            ["KMP (Android + iOS + desktop)", "Yes", "commonMain dependency + platform DSLs"],
            ["Pure iOS (Swift UI only)", "Yes, with Kotlin framework layer", "Thin KMP module + Xcode link (see ios-sample)"],
            ["Pure Swift (zero Kotlin)", "No — not in 1.0", "Deferred; not on roadmap for next release"],
            ["JVM desktop", "Yes", "SyncForge.desktop { }"],
            ["Native macOS", "Yes", "SyncForge.macos { }"],
        ],
    )
    doc.add_paragraph(
        "Consumers still provide: entity storage (Room on Android), backend matching REST_API.md, "
        "and conflict UI for deferToUser() entities."
    )

    # 7. Maven Central checklist
    doc.add_heading("7. Maven Central checklist", level=1)
    doc.add_paragraph("Full details: docs/MAVEN_PUBLISH.md")
    doc.add_heading("One-time setup", level=2)
    bullets(
        doc,
        [
            "Sonatype Central account; namespace dev.syncforge verified.",
            "GitHub secrets: MAVEN_CENTRAL_USERNAME, MAVEN_CENTRAL_PASSWORD, SIGNING_IN_MEMORY_KEY.",
        ],
    )
    doc.add_heading("0.9.0-rc.1 (done / in progress)", level=2)
    bullets(
        doc,
        [
            "verifyReleaseSignOff passed locally.",
            "Tag v0.9.0-rc.1 pushed — Publish Release workflow triggered.",
            "Close + release staging repository in Sonatype Central Portal.",
            "Verify: curl BOM POM from repo1.maven.org.",
            "consumer-smoke/android-minimal compiles with mavenCentral() only.",
        ],
    )
    doc.add_heading("1.0.0 publish", level=2)
    bullets(
        doc,
        [
            "Bump syncforge.version and consumer-smoke pins to 1.0.0.",
            "Update CHANGELOG with 1.0.0 section and API stability notes.",
            "./gradlew verifyReleaseSignOff",
            "git tag v1.0.0 && git push origin v1.0.0",
            "Close + release staging; re-run Central verification.",
        ],
    )
    doc.add_heading("Published artifacts", level=2)
    bullets(
        doc,
        [
            "dev.syncforge:syncforge — main KMP library",
            "dev.syncforge:syncforge-annotations — KSP annotations",
            "dev.syncforge:syncforge-ksp — KSP processor",
            "dev.syncforge:syncforge-persistence — SQLDelight persistence",
            "dev.syncforge:syncforge-android-deps — Room / WorkManager / serialization bundle",
            "dev.syncforge:syncforge-bom — version alignment BOM",
            "dev.syncforge:syncforge-gradle-plugin — id(dev.syncforge.android)",
        ],
    )

    # 8. GitHub & growth playbook
    doc.add_heading("8. GitHub & growth playbook", level=1)
    doc.add_paragraph(
        "SyncForge is a niche library (offline-first sync for mobile/KMP). Growth comes from clear "
        "positioning and lowering integration cost — not from competing on GitHub stars alone."
    )
    doc.add_heading("Positioning (pick one headline)", level=2)
    bullets(
        doc,
        [
            "Room stays source of truth — SyncForge handles the outbox, push/pull, and conflicts.",
            "Offline-first sync for Android without Firebase.",
            "KMP sync with built-in conflict resolution and a runnable backend starter.",
        ],
    )
    doc.add_heading("High-impact actions", level=2)
    add_table(
        doc,
        ["Action", "Why it works"],
        [
            ["60-second demo GIF/video", "Offline edit → airplane mode → sync → conflict resolve — sync is visual"],
            ["Blog post: Sync in 10 minutes", "Link Getting Started; concrete stack names (Room, KSP, WorkManager, SwiftUI)"],
            ["Honest comparison doc", "SyncForge vs roll-your-own vs Firebase vs PowerSync — developers search comparisons"],
            ["Promote :backend-starter", "Library + runnable server in one clone — rare and shareable"],
            ["Update README hero", "Badges: Maven Central, KMP, Apache 2.0; one-liner + quick-start snippet"],
            ["Show HN / r/androiddev / Kotlin Slack", "Launch at 1.0 on Central — RC is too early for Show HN"],
        ],
    )
    doc.add_heading("Lower ROI for this stage", level=2)
    bullets(
        doc,
        [
            "Paid advertising.",
            "Generic #kotlin hashtag spam.",
            "Feature sprawl before anyone uses 1.0.",
            "Chasing star count vs established libraries.",
        ],
    )
    doc.add_heading("Compounding cadence", level=2)
    bullets(
        doc,
        [
            "1.0 — stable + Maven Central + one launch post.",
            "1.0.1 — fix early-adopter friction; Upgrading doc.",
            "Case study — even sample app counts (tasks/notes/tags multi-entity).",
            "Monthly small content (recipe, conflict pattern, E2E tip) beats one launch then silence.",
        ],
    )

    # 9. Launch content kit
    doc.add_heading("9. Launch content kit", level=1)
    doc.add_heading("README improvements (GitHub front page)", level=2)
    bullets(
        doc,
        [
            "Update version badge to 1.0.0 after publish.",
            "Add Maven Central dependency snippet at top (plugin + BOM + syncforge).",
            "Add 30-second value prop: what you get vs what you still provide.",
            "Embed or link demo GIF (offline → sync → conflict).",
            "Link :backend-starter prominently (./gradlew :backend-starter:run).",
            "Platforms row: Android | iOS (KMP) | Desktop | macOS.",
        ],
    )
    doc.add_heading("Demo script (60 seconds)", level=2)
    bullets(
        doc,
        [
            "Terminal 1: ./gradlew :backend-starter:run (or :mock-server:run).",
            "Terminal 2: ./gradlew :sample:installDebug on emulator.",
            "Add task offline → show outbox in debug console (SF button).",
            "Toggle airplane mode off → Sync → task on server.",
            "Create conflict (server edit + local edit) → resolve via conflict chip.",
        ],
    )
    doc.add_heading("Show HN / launch post draft (title options)", level=2)
    bullets(
        doc,
        [
            "Show HN: SyncForge – offline-first sync for Android/KMP with Room as source of truth",
            "Dev.to: Build offline-first Android apps without Firebase in 10 minutes",
            "Tweet/thread: Library + Ktor backend starter + conflict UI — all open source (Apache 2.0)",
        ],
    )
    doc.add_heading("Sample launch post outline", level=2)
    bullets(
        doc,
        [
            "Problem: rolling your own sync (outbox, retry, conflicts) takes months.",
            "Solution: SyncForge — SQLDelight outbox, KSP handlers, conflict strategies, debug console.",
            "Proof: runnable backend-starter + Android Compose sample + iOS SwiftUI sample.",
            "Try it: Maven coordinates + 3-line Android setup.",
            "Ask: feedback on REST API and iOS integration docs.",
        ],
    )

    # 10. Timeline
    doc.add_heading("10. Timeline", level=1)
    add_table(
        doc,
        ["When", "Work"],
        [
            ["Now", "Watch 0.9.0-rc.1 Publish Release; fix CI/secrets if needed"],
            ["Week 1", "Verify Central artifacts; dogfood rc.1; file issues"],
            ["Week 2", "API stability pass; doc version bump; optional auth E2E"],
            ["Ship", "Tag v1.0.0, publish, verify Central, launch post + README refresh"],
            ["Week after 1.0", "1.0.1 if friction reports; start comparison doc / demo GIF"],
        ],
    )

    # 11. Post-1.0 backlog
    doc.add_heading("11. Post-1.0 backlog (1.1+)", level=1)
    bullets(
        doc,
        [
            "DataStore Preferences multiplatform cursor.",
            "SPM / CocoaPods documented consumer path for iOS.",
            "Packaged :sample-desktop app.",
            "Supabase optional transport adapter.",
            "Graduate built-in auth from Experimental to Stable.",
            "Spring backend starter kit.",
            "Swift facade SDK over KMP framework (not full Swift rewrite).",
        ],
    )

    # 12. Decisions not taking
    doc.add_heading("12. Decisions we are not taking for 1.0", level=1)
    bullets(
        doc,
        [
            "Pure Swift library with no Kotlin toolchain.",
            "Large new features (major auth expansion, new conflict models).",
            "Breaking REST contract changes.",
            "Waiting months after rc.1 without a clear blocking reason.",
        ],
    )

    doc.add_paragraph()
    closing = doc.add_paragraph(
        "Next action: confirm Maven Central publish for v0.9.0-rc.1, then begin 1.0 soak. "
        "Regenerate this document after major roadmap changes: "
        "python3 scripts/generate-launch-playbook-docx.py"
    )
    closing.runs[0].italic = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
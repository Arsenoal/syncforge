#!/usr/bin/env python3
"""Generate SyncForge 1.1.0 GitHub issues breakdown DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SyncForge-1.1-Issues.docx"

DOC_DATE = "7 July 2026"
BASELINE = "1.0.0 GA"
TARGET = "1.1.0"


def _style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)


def _shade_cell(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        _shade_cell(hdr[i], "D5E8F0")
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, text in enumerate(row):
            cells[ci].text = text
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _issue_table(doc: Document, rows: list[tuple[str, str, str, str, str]]) -> None:
    _add_table(
        doc,
        ["Issue", "Title", "Job ID", "Depends on", "Done when"],
        rows,
        [0.55, 2.0, 0.55, 0.7, 2.5],
    )


def build() -> None:
    doc = Document()
    _style_doc(doc)

    title = doc.add_heading(f"SyncForge {TARGET} — GitHub Issues Breakdown", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Milestone: Wire-up · Injectable HTTP · EntityStore · Auth hardening")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(60, 60, 60)

    meta = doc.add_paragraph()
    meta.add_run(f"Document date: {DOC_DATE}\n").italic = True
    meta.add_run(f"Baseline: {BASELINE}\n").italic = True
    meta.add_run("Source: docs/ROADMAP_1_0_TO_2_0.md § 1.1.0 GitHub issues breakdown").italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading("Release goal", level=1)
    doc.add_paragraph(
        f"Ship {TARGET} with P0 network and entity-store tracks complete, security/cursor/docs "
        "P0 done, and no breaking changes to 1.0 stable APIs. P1 integration artifacts "
        "(syncforge-integration-*, syncforge-store-realm) may ship in 1.1.1 if needed."
    )

    doc.add_heading("Labels", level=2)
    _bullets(
        doc,
        [
            "Milestone: 1.1.0",
            "Epic labels: epic:network, epic:store, epic:security, epic:dx, epic:release",
            "Area: area:network, area:store, area:security, area:dx, area:docs",
            "Priority: priority:p0, priority:p1, priority:p2",
        ],
    )

    doc.add_heading("Suggested timeline (6 weeks, part-time)", level=2)
    _add_table(
        doc,
        ["Week", "Focus"],
        [
            ("1–2", "A1→A2 (SyncHttpClient, RestSyncTransport) ∥ B1→B2 (EntityStore, @SyncForgeStore)"),
            ("3–4", "A4 httpClient DSL · B3/B4 store adapters · E1 DataStore cursor"),
            ("5", "C1/C2 token security · D1 DI recipes · docs (A5, B5)"),
            ("6", "D2/D3 optional · F1 BOM · acceptance F2 → tag F3"),
        ],
        [0.8, 5.5],
    )

    doc.add_page_break()
    doc.add_heading("Epic A — Injectable HTTP client (P0)", level=1)
    _issue_table(
        doc,
        [
            ("A1", "feat(network): add SyncHttpClient contract in commonMain", "1.1-13", "—", "postPush/getPull API; fake-client unit tests"),
            ("A2", "feat(network): RestSyncTransport delegates to SyncHttpClient", "1.1-14", "A1", "KtorSyncTransport delegates internally; external API unchanged"),
            ("A3", "refactor(network): extract :syncforge-network-ktor adapter", "1.1-15", "A2", "Ktor wiring in optional module"),
            ("A4", "feat(dsl): httpClient { } on platform DSLs", "1.1-16", "A2", "Injected app HttpClient; bundled fallback"),
            ("A5", "docs(network): injectable HttpClient guide", "1.1-12", "A4", "GETTING_STARTED + RECIPES interceptors sample"),
        ],
    )

    doc.add_heading("Epic B — EntityStore abstraction (P0)", level=1)
    _issue_table(
        doc,
        [
            ("B1", "feat(store): EntityStore contract in commonMain", "1.1-09", "—", "CRUD + transaction; maps to EntitySyncHandler"),
            ("B2", "feat(ksp): @SyncForgeStore handler generation", "1.1-10", "B1", "@SyncForgeDao path unchanged"),
            ("B3", "feat(store): :syncforge-store-room adapter", "1.1-11", "B2", "Optional Maven artifact; not in core BOM transitives"),
            ("B4", "feat(store): in-memory EntityStore for commonTest", "1.1-11", "B1", "Non-Room path without Realm"),
            ("B5", "docs(store): BYO store path in GETTING_STARTED", "1.1-12", "B2, B4", "Room optional; plugin skips Room KSP when unused"),
        ],
    )

    doc.add_heading("Epic C — Auth & token security (P1, GA-required)", level=1)
    _issue_table(
        doc,
        [
            ("C1", "feat(auth): encrypted TokenStore (Android + iOS Keychain)", "1.1-17", "—", "Migration from plain prefs documented"),
            ("C2", "feat(auth): login/register CharArray overloads", "1.1-18", "—", "Wipe in finally; AUTH_API updated"),
            ("C3", "api(auth): graduate built-in auth DSL to stable", "1.1-04", "C1, C2", "Remove @ExperimentalSyncForgeApi on auth surfaces"),
        ],
    )

    doc.add_heading("Epic D — DI & developer experience", level=1)
    _issue_table(
        doc,
        [
            ("D1", "docs(dx): Koin + Hilt recipes in RECIPES.md", "1.1-01", "B1", "No Koin/Dagger in :syncforge core"),
            ("D2", "feat(dx): publish syncforge-integration-koin", "1.1-02", "D1", "syncForgeModule { } + WorkManager helper"),
            ("D3", "feat(dx): publish syncforge-integration-hilt", "1.1-03", "D1", "@Provides templates"),
            ("D4", "feat(sample): optional :sample-di variant", "1.1-06", "D2 or D3", "Documented DI wiring"),
        ],
    )

    doc.add_heading("Epic E — Cursor persistence (P1)", level=1)
    _issue_table(
        doc,
        [
            ("E1", "feat(persistence): DataStore Preferences pull cursor (Android)", "1.1-05", "—", "iOS UserDefaults fallback documented"),
        ],
    )

    doc.add_page_break()
    doc.add_heading("Epic F — 1.1.0 release gate", level=1)
    _issue_table(
        doc,
        [
            ("F1", "chore(dist): BOM lists optional 1.1 artifacts", "acceptance", "A3, B3", "Store/network/integration in BOM constraints only"),
            ("F2", "test(qa): 1.1.0 acceptance matrix", "§ criteria", "A*, B*, C*, D1, E1", "All acceptance checkboxes green"),
            ("F3", "chore(release): tag v1.1.0 + Maven Central publish", "1.1-08", "F2", "CHANGELOG; verify workflow"),
            ("F4", "docs: MODULES + CHANGELOG 1.1.0 freeze", "—", "F2", "Stability table updated"),
        ],
    )

    doc.add_heading("1.1.0 acceptance criteria", level=1)
    _bullets(
        doc,
        [
            "SyncHttpClient + RestSyncTransport published; KtorSyncTransport backward compatible",
            "Injectable Ktor HttpClient documented with app-owned client sample",
            "EntityStore in commonMain; EntitySyncHandler delegates through it",
            "KSP generates handlers from @SyncForgeStore and existing @SyncForgeDao",
            "Non-Room path documented (in-memory store test or recipe)",
            "RECIPES.md DI section — Koin + Hilt examples",
            "DataStore cursor on Android; iOS fallback documented",
            "Encrypted TokenStore default; CharArray login/register overloads",
            "BOM lists optional integration + store artifacts (not transitive)",
            "No breaking changes to 1.0 stable APIs",
        ],
    )

    doc.add_heading("1.0.x patch lane", level=1)
    doc.add_paragraph(
        "Job 1.1-08: semver patches on 1.0.x branch only — bugfixes, no new APIs. "
        "main targets 1.1.0 feature work."
    )

    doc.add_heading("Related documents", level=1)
    _add_table(
        doc,
        ["Document", "Purpose"],
        [
            ("docs/ROADMAP_1_0_TO_2_0.md", "Markdown source — full 1.1.x architecture"),
            ("docs/ROADMAP.md", "High-level next: 1.1.0 summary"),
            ("docs/SyncForge-Roadmap-1.0-to-2.0.docx", "Full roadmap Word export"),
        ],
        [2.8, 3.5],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
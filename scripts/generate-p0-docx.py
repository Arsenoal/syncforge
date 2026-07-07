#!/usr/bin/env python3
"""Generate short P0 checklist DOCX for SyncForge 1.0.0."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SyncForge-1.0-P0.docx"

# Remaining work before tagging v1.0.0 (completed items removed July 2026).
SECTIONS = [
    (
        "Distribution (1.0-P0-03)",
        [
            "Publish 1.0.0 to Maven Central — tag v1.0.0; all artifacts (syncforge, annotations, ksp, persistence, android-deps, bom, gradle-plugin)",
        ],
    ),
    (
        "Documentation (P1-01)",
        [
            "Upgrade guide — pre-0.6 Room → SQLDelight (recommended; migration path in ANDROID_SETUP.md)",
        ],
    ),
    (
        "CI & QA (1.0-P0-05, P1-03)",
        [
            "macOS tag publish — iOS/macOS frameworks from publish-release.yml without manual steps",
            "ProGuard/R8 — document and test consumer-rules.pro (P1-03)",
        ],
    ),
    (
        "Community soak (P1-07)",
        [
            "At least one external dogfood or documented third-party integration (#feed thread, GitHub issues)",
        ],
    ),
    (
        "Release gate (post-soak)",
        [
            "Re-run sign-off matrix at v1.0.0 — ./gradlew verifySignOffMatrix + CI E2E + publish-release.yml on tag",
            "Bump consumer-smoke Maven Central pins to 1.0.0 after Central sync",
            "Tag v1.0.0",
        ],
    ),
]


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("SyncForge 1.0.0 — P0 Checklist", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Remaining work before first stable release (1.0.0).")

    work_items = sum(len(items) for heading, items in SECTIONS if "sign-off" not in heading.lower())
    sign_off_items = sum(
        len(items) for heading, items in SECTIONS if "sign-off" in heading.lower()
    )

    doc.add_paragraph("Date: July 2026")
    doc.add_paragraph("Current baseline: 0.9.0-rc.5 (Maven Central complete)")
    doc.add_paragraph(f"Remaining work items: {work_items}")
    doc.add_paragraph(f"Release sign-off checks: {sign_off_items}")
    doc.add_paragraph()

    doc.add_heading("Release gate", level=1)
    doc.add_paragraph(
        "Tag 1.0.0 only when every work item below is complete and every sign-off check passes. "
        "Maven publish: docs/MAVEN_PUBLISH.md."
    )

    for heading, items in SECTIONS:
        doc.add_heading(heading, level=2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT} ({work_items} work items, {sign_off_items} sign-off checks)")


if __name__ == "__main__":
    build()